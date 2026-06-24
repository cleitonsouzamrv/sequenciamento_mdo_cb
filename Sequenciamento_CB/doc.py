from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ──────────────────────────────────────────────
# HELPERS
# ──────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Pinta o fundo de uma célula com a cor hex informada (ex: '006B3F')."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_header_row(table, headers: list, bg_hex='006B3F', font_color='FFFFFF'):
    """Preenche a primeira linha da tabela como cabeçalho estilizado."""
    row = table.rows[0]
    for i, text in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.bold           = True
        run.font.color.rgb = RGBColor.from_string(font_color)
        run.font.size      = Pt(10)
        set_cell_bg(cell, bg_hex)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_data_rows(table, data: list, start_row=1, alt_color='E6F4EE'):
    """Preenche linhas de dados com zebra-striping."""
    for i, row_data in enumerate(data):
        row = table.rows[start_row + i]
        for j, text in enumerate(row_data):
            cell      = row.cells[j]
            cell.text = str(text)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if i % 2 == 1:
                set_cell_bg(cell, alt_color)


def add_table(doc, headers: list, data: list, col_widths: list = None):
    """Cria tabela completa com cabeçalho + dados."""
    table       = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Table Grid'
    add_header_row(table, headers)
    add_data_rows(table, data)
    if col_widths:
        for row in table.rows:
            for j, width in enumerate(col_widths):
                row.cells[j].width = Inches(width)
    return table


def add_bullet(doc, text: str, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def add_numbered(doc, text: str):
    return doc.add_paragraph(text, style='List Number')


def section_divider(doc):
    doc.add_paragraph('')


# ──────────────────────────────────────────────
# DOCUMENTO
# ──────────────────────────────────────────────

doc  = Document()
base = doc.styles['Normal']
base.font.name = 'Calibri'
base.font.size = Pt(11)

# ── CAPA ──────────────────────────────────────
titulo = doc.add_heading('Documentação Técnica', level=0)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitulo = doc.add_paragraph('Sequenciador de Obras — Capacete Branco MRV')
subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitulo.runs[0].bold           = True
subtitulo.runs[0].font.size      = Pt(14)
subtitulo.runs[0].font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)

data_doc = doc.add_paragraph(f'Versão 2.1  ·  Gerada em {datetime.date.today().strftime("%d/%m/%Y")}')
data_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
data_doc.runs[0].font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)

doc.add_page_break()

# ══════════════════════════════════════════════
# SUMÁRIO
# ══════════════════════════════════════════════
doc.add_heading('Sumário', level=1)

sumario_itens = [
    '1.  Visão Geral',
    '2.  Objetivo',
    '3.  Fluxo Geral da Aplicação',
    '4.  Planilhas de Entrada',
    '    4.1  Guia Base — Estrutura esperada',
    '    4.2  Guia Sheet1 — Estrutura esperada',
    '5.  Configurações do Usuário',
    '    5.1  Marco de Início da Obra A',
    '    5.2  Filtro de Latência Máxima',
    '6.  Pré-processamento',
    '    6.1  Montagem da Linha (pivot)',
    '    6.2  Resolução de Simultaneidade',
    '    6.3  Filtro de Obras Candidatas',
    '7.  Critérios de Sequenciamento — Linhas Existentes',
    '    7.1  1ª Tentativa: Mesmo Cluster',
    '    7.2  2ª Tentativa: Fallback para Outros Clusters',
    '8.  Criação de Novas Linhas',
    '    8.1  O que é uma Nova Linha',
    '    8.2  Quando uma Nova Linha é criada',
    '    8.3  Fluxo de decisão',
    '    8.4  Nomenclatura',
    '9.  Planilha de Saída — Colunas Geradas',
    '10. Regras de Complexidade por Senioridade',
    '11. Geocodificação e Distâncias',
    '12. Tecnologias Utilizadas',
    '13. Limitações e Observações',
]
for item in sumario_itens:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.25 if item.startswith('    ') else 0)
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════
# 1. VISÃO GERAL
# ══════════════════════════════════════════════
doc.add_heading('1. Visão Geral', level=1)
doc.add_paragraph(
    'O Sequenciador de Capacete Branco é uma aplicação web desenvolvida em Python com Streamlit '
    'que automatiza a alocação da próxima obra para colaboradores (Capacetes Brancos) da MRV. '
    'A ferramenta lê duas planilhas Excel, processa as obras atuais de cada colaborador e sugere '
    'a próxima obra com base em critérios de cluster, janela temporal, complexidade e proximidade geográfica. '
    'Obras que não se encaixam em nenhuma linha existente são agrupadas em novas linhas criadas automaticamente.'
)
section_divider(doc)

# ══════════════════════════════════════════════
# 2. OBJETIVO
# ══════════════════════════════════════════════
doc.add_heading('2. Objetivo', level=1)
doc.add_paragraph(
    'Garantir que cada Capacete Branco, ao finalizar sua obra atual, tenha uma próxima obra '
    'sequenciada de forma otimizada, minimizando o tempo de ociosidade (latência) e a distância '
    'de deslocamento, respeitando a estrutura de clusters, a senioridade do profissional e a '
    'complexidade das obras.'
)
section_divider(doc)

# ══════════════════════════════════════════════
# 3. FLUXO GERAL
# ══════════════════════════════════════════════
doc.add_heading('3. Fluxo Geral da Aplicação', level=1)

etapas = [
    'Upload das duas planilhas .xlsx pelo usuário (Planilha Base + Todos Empreendimentos).',
    'Validação das colunas obrigatórias em cada guia.',
    'Conversão e normalização das datas.',
    'Montagem da coluna "Linha" (pivot: CLUSTER | Sigla Senioridade | Responsável 1).',
    'Resolução de simultaneidade: obras simultâneas na mesma linha são devolvidas ao pool.',
    'Geocodificação das cidades (coordenadas via Nominatim/OpenStreetMap).',
    'Sequenciamento das linhas existentes: 1ª tentativa no mesmo cluster; se não encontrar, '
    'fallback para outros clusters dentro do raio de 200 km e da latência máxima configurada.',
    'Criação de novas linhas: obras não alocadas são agrupadas em novas linhas por cluster.',
    'Montagem do output empilhado por linha e ordem.',
    'Exibição do resultado na tela + download da planilha Excel gerada.',
]
for etapa in etapas:
    add_numbered(doc, etapa)
section_divider(doc)

# ══════════════════════════════════════════════
# 4. PLANILHAS DE ENTRADA
# ══════════════════════════════════════════════
doc.add_heading('4. Planilhas de Entrada', level=1)

# 4.1 Base
doc.add_heading('4.1  Guia Base', level=2)
doc.add_paragraph(
    'Planilha principal com os engenheiros/responsáveis e suas obras atuais. '
    'O app localiza automaticamente a linha de cabeçalho buscando a coluna "OBRA".'
)

headers_base = ['Coluna', 'Tipo', 'Obrigatório', 'Descrição']
data_base = [
    ('OBRA',                          'Texto', '✅', 'Nome do empreendimento/obra'),
    ('Cód. CRM',                      'Texto', '✅', 'Código identificador da obra no CRM'),
    ('Regional',                      'Texto', '✅', 'Regional responsável pela obra'),
    ('Cidade',                        'Texto', '✅', 'Cidade onde a obra está localizada'),
    ('CLUSTER',                       'Texto', '✅', 'Cluster ao qual a obra pertence'),
    ('AMP (abr)',                      'Texto', '❌', 'Abreviação do AMP da obra'),
    ('Status',                        'Texto', '✅', 'Status atual da obra'),
    ('Data Fundação',                 'Data',  '✅', 'Data de fundação/início da obra'),
    ('Data Terraplenagem',            'Data',  '✅', 'Marco de início para o sequenciamento'),
    ('Data Encerramento Módulo',      'Data',  '✅', 'Marco de término para o sequenciamento'),
    ('COMPLEXIDADE',                  'Texto', '❌', 'Nível de complexidade (NÃO COMPLEXO / MÉDIO / COMPLEXO)'),
    ('Engenheiro 1',                  'Texto', '✅', 'Nome do primeiro engenheiro responsável'),
    ('Engenheiro 2',                  'Texto', '❌', 'Nome do segundo engenheiro (quando houver)'),
    ('Engenheiro CRM',                'Texto', '❌', 'Engenheiro registrado no CRM'),
    ('Gestor CRM',                    'Texto', '❌', 'Gestor registrado no CRM'),
    ('Responsável 1',                 'Texto', '✅', 'Responsável principal pela obra'),
    ('Responsável 2 (quando houver)', 'Texto', '❌', 'Segundo responsável (quando houver)'),
    ('Senioridade ENG1',              'Texto', '❌', 'Nível de senioridade do Engenheiro 1'),
    ('Senioridade ENG2',              'Texto', '❌', 'Nível de senioridade do Engenheiro 2'),
    ('Gestor',                        'Texto', '❌', 'Gestor responsável pela obra'),
    ('Observação',                    'Texto', '❌', 'Observações gerais sobre a obra'),
]
add_table(doc, headers_base, data_base, col_widths=[2.0, 0.7, 0.9, 3.2])
section_divider(doc)

# 4.2 Sheet1
doc.add_heading('4.2  Guia Sheet1 — Todos Empreendimentos', level=2)
doc.add_paragraph(
    'Catálogo completo de empreendimentos disponíveis para sequenciamento. '
    'O app localiza o cabeçalho buscando a coluna "IDTRSICRM".'
)

headers_emp = ['Coluna', 'Tipo', 'Obrigatório', 'Descrição']
data_emp = [
    ('Regional',            'Texto', '✅', 'Regional responsável pelo empreendimento'),
    ('IDTRSICRM',           'Texto', '✅', 'Identificador único do empreendimento no CRM'),
    ('Empreendimento',      'Texto', '✅', 'Nome do empreendimento'),
    ('CLUSTER_CORRIGIDO',   'Texto', '✅', 'Cluster corrigido do empreendimento'),
    ('CIDADE',              'Texto', '✅', 'Cidade onde o empreendimento está localizado'),
    ('Fundação',            'Data',  '✅', 'Data prevista de início da fundação'),
    ('Terraplenagem',       'Data',  '✅', 'Marco de início para o sequenciamento'),
    ('Encerramento Módulo', 'Data',  '✅', 'Data prevista de encerramento do módulo'),
    ('Complexidade',        'Texto', '❌', 'Nível de complexidade da obra candidata'),
    ('Nome Ref Ajuste',     'Texto', '❌', 'Nome de referência para ajuste — informativo'),
]
add_table(doc, headers_emp, data_emp, col_widths=[1.8, 0.7, 0.9, 3.3])
section_divider(doc)

# ══════════════════════════════════════════════
# 5. CONFIGURAÇÕES DO USUÁRIO
# ══════════════════════════════════════════════
doc.add_heading('5. Configurações do Usuário', level=1)

# 5.1
doc.add_heading('5.1  Marco de Início da Obra A', level=2)
doc.add_paragraph(
    'Define qual data será registrada na coluna "Marco Início Obra A" '
    'para as linhas existentes (Ordem 1).'
)

headers_cfg = ['Opção', 'Coluna utilizada', 'Quando usar']
data_cfg = [
    ('Terraplenagem', 'Data Terraplenagem (guia Base)', 'Padrão — marco de início operacional da obra'),
    ('Fundação',      'Data Fundação (guia Base)',      'Quando o marco relevante for o início da fundação'),
]
add_table(doc, headers_cfg, data_cfg, col_widths=[1.5, 2.5, 3.2])
section_divider(doc)

# 5.2
doc.add_heading('5.2  Filtro de Latência Máxima', level=2)
doc.add_paragraph(
    'O usuário define a latência máxima permitida entre o encerramento de uma obra e o início '
    '(terraplenagem) da próxima. Esse filtro é aplicado tanto no sequenciamento de linhas '
    'existentes quanto na criação de novas linhas.'
)
add_bullet(doc, 'Intervalo permitido: de 3 meses a 7 meses.')
add_bullet(doc, 'Valor padrão: 3 meses (90 dias).')
add_bullet(doc, 'Internamente o valor em meses é convertido para dias (meses × 30).')
add_bullet(doc,
    'Obras cuja latência calculada exceder o limite configurado são descartadas '
    'da tentativa atual e podem originar uma nova linha.'
)
section_divider(doc)

# ══════════════════════════════════════════════
# 6. PRÉ-PROCESSAMENTO
# ══════════════════════════════════════════════
doc.add_heading('6. Pré-processamento', level=1)

doc.add_heading('6.1  Montagem da Linha (pivot)', level=2)
doc.add_paragraph(
    'A coluna "Linha" é o identificador principal de cada sequência de obras. '
    'Ela é construída automaticamente pelo app concatenando três campos da guia Base:'
)
add_bullet(doc, 'CLUSTER  →  ex.: "SP-01"')
add_bullet(doc, 'Sigla de Senioridade  →  "EN" (Engenheiro/a) ou "AN" (Analista)')
add_bullet(doc, 'Responsável 1  →  nome em maiúsculas')
doc.add_paragraph('')
doc.add_paragraph('Exemplo de Linha gerada:  SP-01 | EN | JOÃO SILVA').runs[0].italic = True

doc.add_heading('6.2  Resolução de Simultaneidade', level=2)
doc.add_paragraph(
    'Quando uma mesma Linha possui mais de uma obra na guia Base (obras simultâneas), '
    'o app mantém apenas a obra com a maior Data de Encerramento Módulo como representante ativa. '
    'As demais são marcadas com "Sim — devolvida ao pool" na coluna Simultaneidade e '
    'reinseridas no pool de obras não alocadas para criação de novas linhas.'
)

doc.add_heading('6.3  Filtro de Obras Candidatas', level=2)
doc.add_paragraph('Antes do sequenciamento, o app filtra a guia Sheet1 mantendo apenas obras que:')
add_bullet(doc, 'Possuem Data Terraplenagem, Data Encerramento Módulo e Cidade preenchidas.')
add_bullet(doc, 'Têm Data Terraplenagem ≥ data de hoje (obras futuras).')
add_bullet(doc, 'Não estão presentes na guia Base (não são obras já em andamento).')
section_divider(doc)

# ══════════════════════════════════════════════
# 7. CRITÉRIOS DE SEQUENCIAMENTO — LINHAS EXISTENTES
# ══════════════════════════════════════════════
doc.add_heading('7. Critérios de Sequenciamento — Linhas Existentes', level=1)
doc.add_paragraph(
    'Para cada linha existente (guia Base) que ainda não possui "Próxima Obra" definida, '
    'o app realiza duas tentativas de sequenciamento antes de desistir.'
)

# 7.1
doc.add_heading('7.1  1ª Tentativa: Mesmo Cluster', level=2)
doc.add_paragraph(
    'O app busca candidatas dentro do mesmo cluster da linha existente, '
    'aplicando os seguintes filtros em ordem:'
)

headers_crit = ['#', 'Critério', 'Regra']
data_crit = [
    ('1', 'Janela Temporal',  'Data Terraplenagem da candidata ≥ Data Encerramento Módulo da obra atual'),
    ('2', 'Mesmo Cluster',    'CLUSTER_CORRIGIDO da candidata = CLUSTER da linha existente'),
    ('3', 'Complexidade',     'Compatível com a senioridade do Engenheiro 1 (ver seção 10)'),
    ('4', 'Distância',        'Distância entre cidades ≤ 200 km (via Geopy/Nominatim)'),
    ('5', 'Latência Máxima',  'Gap em dias entre encerramento e terraplenagem ≤ latência máxima configurada'),
    ('6', 'Disponibilidade',  'Obra ainda não alocada em nenhuma outra linha'),
    ('7', 'Menor gap (dias)', 'Dentre as candidatas válidas, seleciona a de menor gap em dias'),
]
add_table(doc, headers_crit, data_crit, col_widths=[0.4, 1.8, 5.0])
section_divider(doc)

# 7.2
doc.add_heading('7.2  2ª Tentativa: Fallback para Outros Clusters', level=2)
doc.add_paragraph(
    'Se nenhuma candidata for encontrada no mesmo cluster, o app realiza uma segunda busca '
    'considerando obras de outros clusters. Os mesmos filtros de complexidade, distância (≤ 200 km) '
    'e latência máxima são aplicados. A obra selecionada recebe a marcação '
    '"Sequenciamento - ferramenta (cluster diferente)" na coluna Origem Sequenciamento, '
    'permitindo ao usuário identificar facilmente os casos de cruzamento de cluster.'
)

headers_fb = ['Situação', 'Origem Sequenciamento gravada']
data_fb = [
    ('Sequenciou no mesmo cluster',     'Sequenciamento - ferramenta'),
    ('Sequenciou em cluster diferente', 'Sequenciamento - ferramenta (cluster diferente)'),
    ('Veio preenchido na guia Base',    'Sequenciamento DH'),
    ('Obra em nova linha criada',       'Sequenciamento - ferramenta'),
]
add_table(doc, headers_fb, data_fb, col_widths=[3.5, 4.0])
section_divider(doc)

# ══════════════════════════════════════════════
# 8. CRIAÇÃO DE NOVAS LINHAS
# ══════════════════════════════════════════════
doc.add_heading('8. Criação de Novas Linhas', level=1)

# 8.1
doc.add_heading('8.1  O que é uma Nova Linha', level=2)
doc.add_paragraph(
    'Uma Nova Linha representa uma sequência de obras que não está vinculada a nenhum '
    'colaborador existente na guia Base. Ela é criada pelo app para acomodar obras do '
    'catálogo (Sheet1) que não puderam ser alocadas como "próxima obra" de nenhuma linha existente. '
    'Na saída, essas linhas aparecem com Tipo de Linha = "Nova" e recebem o nome no formato:'
)
p = doc.add_paragraph('{CLUSTER} | Linha nova {N}')
p.runs[0].italic = True
p.runs[0].bold   = True
p.paragraph_format.left_indent = Inches(0.5)
doc.add_paragraph('Exemplo:  SP-01 | Linha nova 1').runs[0].italic = True

# 8.2
doc.add_heading('8.2  Quando uma Nova Linha é criada', level=2)
doc.add_paragraph(
    'O app tenta encaixar cada obra não alocada em duas etapas antes de criar uma nova linha:'
)

doc.add_paragraph('Etapa 1 — Linhas Existentes (guia Base)').runs[0].bold = True
doc.add_paragraph(
    'O app tenta alocar a obra como "Próxima Obra" de uma linha existente, '
    'primeiro no mesmo cluster e depois em outros clusters (fallback). '
    'Se nenhuma linha existente aceitar a obra, ela vai para o pool de não alocadas.'
)

doc.add_paragraph('Etapa 2 — Novas Linhas já criadas nesta execução').runs[0].bold = True
doc.add_paragraph(
    'O app verifica se alguma nova linha já criada pode receber a obra. '
    'Os critérios são: mesmo cluster, terraplenagem ≥ encerramento da última obra da linha, '
    'distância ≤ 200 km da última cidade e gap ≤ latência máxima configurada. '
    'Dentre as elegíveis, escolhe a de menor gap em dias.'
)

doc.add_paragraph('Criação de Nova Linha').runs[0].bold = True
doc.add_paragraph(
    'Se a obra não se encaixar em nenhuma linha existente nem em nenhuma nova linha já criada, '
    'o app cria uma nova linha para o cluster da obra e a aloca como primeira obra dessa linha.'
)

# 8.3
doc.add_heading('8.3  Fluxo de Decisão', level=2)

fluxo = [
    ('Obra candidata não alocada',                          '↓'),
    ('Cabe em linha EXISTENTE — mesmo cluster?',            'SIM → aloca como Ordem 2 | Origem: Sequenciamento - ferramenta'),
    ('',                                                    'NÃO ↓'),
    ('Cabe em linha EXISTENTE — outro cluster?',            'SIM → aloca como Ordem 2 | Origem: Sequenciamento - ferramenta (cluster diferente)'),
    ('',                                                    'NÃO ↓'),
    ('Cabe em NOVA LINHA já criada (mesmo cluster)?',       'SIM → aloca na nova linha (Ordem N)'),
    ('',                                                    'NÃO ↓'),
    ('🆕 Cria nova linha: {CLUSTER} | Linha nova {N}',      'Obra alocada como Ordem 1 da nova linha'),
]
headers_fluxo = ['Decisão', 'Resultado']
add_table(doc, headers_fluxo, fluxo, col_widths=[3.5, 4.0])

# 8.4
doc.add_heading('8.4  Nomenclatura e Ordenação', level=2)
add_bullet(doc, 'As obras não alocadas são ordenadas por Data Terraplenagem (crescente) antes do agrupamento.')
add_bullet(doc, 'O contador de novas linhas é por cluster: cada cluster tem seu próprio contador (Linha nova 1, Linha nova 2, ...).')
add_bullet(doc, 'Obras devolvidas ao pool por simultaneidade também passam por esse processo e podem originar novas linhas.')
add_bullet(doc, 'A coluna "Origem Sequenciamento" de todas as obras em novas linhas recebe o valor "Sequenciamento - ferramenta".')
section_divider(doc)

# ══════════════════════════════════════════════
# 9. PLANILHA DE SAÍDA
# ══════════════════════════════════════════════
doc.add_heading('9. Planilha de Saída — Colunas Geradas', level=1)
doc.add_paragraph(
    'O resultado é uma tabela empilhada onde cada linha representa uma obra dentro de uma sequência. '
    'As colunas abaixo são geradas automaticamente pelo app:'
)

headers_saida = ['Coluna', 'Tipo', 'Descrição']
data_saida = [
    ('Linha',                    'Texto',  'Pivot principal: CLUSTER | Sigla Senioridade | Responsável 1'),
    ('Tipo de Linha',            'Texto',  '"Existente" = veio da guia Base | "Nova" = criada pelo app'),
    ('Ordem',                    'Número', 'Sequência da obra dentro da linha (1 = atual, 2 = próxima, ...)'),
    ('OBRA',                     'Texto',  'Nome da obra'),
    ('Cód. CRM',                 'Texto',  'Código CRM da obra'),
    ('Regional',                 'Texto',  'Regional da obra'),
    ('Cidade',                   'Texto',  'Cidade da obra'),
    ('CLUSTER',                  'Texto',  'Cluster da obra'),
    ('Complexidade Obra',        'Texto',  'Nível de complexidade da obra'),
    ('Senioridade ENG1',         'Texto',  'Senioridade do Engenheiro 1 (apenas Ordem 1 de linhas existentes)'),
    ('Data Fundação',            'Data',   'Data de fundação da obra'),
    ('Data Terraplenagem',       'Data',   'Data de terraplenagem — marco de início'),
    ('Data Encerramento Módulo', 'Data',   'Data de encerramento — marco de término'),
    ('Marco Início Obra A',      'Data',   'Data usada como marco de início da Obra A (conforme seleção do usuário)'),
    ('Simultaneidade',           'Texto',  '"Sim — devolvida ao pool" = obra estava simultânea e foi redistribuída'),
    ('Fonte da Próxima Obra',    'Texto',  '"DH" = veio da guia Base | "Ferramenta" = sugerida pelo app'),
    ('Origem Sequenciamento',    'Texto',
     '"Sequenciamento DH" = já existia na Base | '
     '"Sequenciamento - ferramenta" = app sequenciou no mesmo cluster | '
     '"Sequenciamento - ferramenta (cluster diferente)" = app sequenciou via fallback de cluster'),
    ('Latência (Dias)',          'Número', 'Gap em dias entre encerramento da obra anterior e terraplenagem desta'),
    ('Distância (km)',           'Número', 'Distância em km da obra anterior para esta'),
]
add_table(doc, headers_saida, data_saida, col_widths=[2.0, 0.8, 4.4])
section_divider(doc)

# ══════════════════════════════════════════════
# 10. REGRAS DE COMPLEXIDADE
# ══════════════════════════════════════════════
doc.add_heading('10. Regras de Complexidade por Senioridade', level=1)
doc.add_paragraph(
    'O app verifica a senioridade do Engenheiro 1 da linha e restringe as obras candidatas '
    'às complexidades permitidas para aquele nível. A regra se aplica tanto na 1ª tentativa '
    '(mesmo cluster) quanto no fallback (outros clusters):'
)

headers_comp = ['Senioridade', 'Sigla gerada', 'Complexidades permitidas']
data_comp = [
    ('Engenheiro(a) Júnior',  'EN', 'NÃO COMPLEXO'),
    ('Engenheiro(a) Pleno',   'EN', 'NÃO COMPLEXO, MÉDIO'),
    ('Engenheiro(a) Sênior',  'EN', 'NÃO COMPLEXO, MÉDIO, COMPLEXO'),
    ('Analista (qualquer)',    'AN', 'Sem restrição de complexidade aplicada'),
    ('Não informado / vazio', '—',  'Sem restrição de complexidade aplicada'),
]
add_table(doc, headers_comp, data_comp, col_widths=[2.2, 1.2, 3.8])
section_divider(doc)

# ══════════════════════════════════════════════
# 11. GEOCODIFICAÇÃO E DISTÂNCIAS
# ══════════════════════════════════════════════
doc.add_heading('11. Geocodificação e Distâncias', level=1)
doc.add_paragraph(
    'O app utiliza o serviço Nominatim (OpenStreetMap) via biblioteca Geopy para obter as '
    'coordenadas geográficas (latitude/longitude) de cada cidade presente nas planilhas. '
    'As coordenadas são armazenadas em cache durante a execução para evitar requisições repetidas. '
    'A distância entre duas cidades é calculada como distância geodésica (linha reta sobre a superfície terrestre). '
    'Obras cuja cidade não for encontrada pelo geocodificador recebem distância 9999 km e são descartadas.'
)
add_bullet(doc, 'Intervalo entre requisições: ~1,1 segundo (respeito ao rate limit do Nominatim).')
add_bullet(doc, 'Limite de distância aplicado: 200 km — válido tanto para mesmo cluster quanto para fallback.')
add_bullet(doc, 'Cidades não encontradas: distância definida como 9999 km → obra descartada.')
section_divider(doc)

# ══════════════════════════════════════════════
# 12. TECNOLOGIAS UTILIZADAS
# ══════════════════════════════════════════════
doc.add_heading('12. Tecnologias Utilizadas', level=1)

headers_tech = ['Tecnologia', 'Versão mínima', 'Uso']
data_tech = [
    ('Python',      '3.9+',   'Linguagem principal'),
    ('Streamlit',   '1.30+',  'Interface web interativa'),
    ('Pandas',      '2.0+',   'Manipulação e transformação de dados'),
    ('Geopy',       '2.3+',   'Geocodificação e cálculo de distâncias geodésicas'),
    ('XlsxWriter',  '3.0+',   'Geração do arquivo Excel de saída'),
    ('python-docx', '1.1+',   'Geração desta documentação'),
    ('unicodedata', 'stdlib', 'Normalização de strings (remoção de acentos)'),
]
add_table(doc, headers_tech, data_tech, col_widths=[1.8, 1.2, 4.2])
section_divider(doc)

# ══════════════════════════════════════════════
# 13. LIMITAÇÕES E OBSERVAÇÕES
# ══════════════════════════════════════════════
doc.add_heading('13. Limitações e Observações', level=1)

headers_lim = ['Item', 'Detalhe']
data_lim = [
    ('Geocodificação',
     'Usa o serviço gratuito Nominatim. Pode haver lentidão (~1s/cidade) e eventual imprecisão para cidades com nomes ambíguos.'),
    ('Alocação sequencial',
     'O algoritmo percorre as linhas na ordem em que aparecem na planilha. A ordem das linhas pode influenciar o resultado final.'),
    ('Sem reotimização global',
     'Não realiza otimização combinatória. Cada linha é resolvida individualmente, sem considerar o impacto nas demais.'),
    ('Linhas já preenchidas',
     'Linhas com "Próxima Obra" já preenchida na guia Base não são alteradas, mas seus IDs são marcados como alocados.'),
    ('Novas linhas sem eng.',
     'Novas linhas criadas pelo app não possuem engenheiro associado — a coluna Senioridade ENG1 fica vazia para essas linhas.'),
    ('Simultaneidade',
     'Apenas a obra de maior encerramento é mantida por linha. As demais são devolvidas ao pool independentemente de outros critérios.'),
    ('Cache de coordenadas',
     'O cache de coordenadas é válido apenas durante a sessão. Ao recarregar a página, as coordenadas são buscadas novamente.'),
    ('Fallback de cluster',
     'O cruzamento de cluster é identificado na coluna "Origem Sequenciamento" com o valor '
     '"Sequenciamento - ferramenta (cluster diferente)". '
     'Recomenda-se revisar esses casos manualmente antes de publicar o sequenciamento.'),
    ('Latência máxima',
     'O filtro de latência (3 a 7 meses) é aplicado tanto no sequenciamento de linhas existentes '
     'quanto na criação de novas linhas. Obras fora da janela não são descartadas permanentemente — '
     'podem originar uma nova linha própria.'),
]
add_table(doc, headers_lim, data_lim, col_widths=[2.0, 5.2])
section_divider(doc)

# ── RODAPÉ ────────────────────────────────────
doc.add_paragraph('')
rodape = doc.add_paragraph(
    f'Documentação gerada automaticamente em {datetime.date.today().strftime("%d/%m/%Y")} '
    '— Sequenciador CB MRV v2.1'
)
rodape.alignment              = WD_ALIGN_PARAGRAPH.CENTER
rodape.runs[0].italic         = True
rodape.runs[0].font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)

# ── SALVAR ────────────────────────────────────
nome_arquivo = 'Documentacao_Sequenciador_CB_MRV_v2.1.docx'
doc.save(nome_arquivo)
print(f'✅ Arquivo gerado: {nome_arquivo}')
